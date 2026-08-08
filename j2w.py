import os
import re
import base64
import docx
from docx.shared import Inches, Pt
import streamlit as st

try:
    from mistralai.client import Mistral
    MISTRAL_AVAILABLE = True
except ImportError:
    MISTRAL_AVAILABLE = False

st.set_page_config(page_title="Convert PDF to Word with python-docx", page_icon="📐", layout="wide")

st.title("📐 Convert PDF to Word (Mistral OCR + python-docx + Preview)")

# Khởi tạo các giá trị trong Session State
if "preview_markdown" not in st.session_state:
    st.session_state.preview_markdown = ""
if "images_dict" not in st.session_state:
    st.session_state.images_dict = {}
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "file_name" not in st.session_state:
    st.session_state.file_name = "Document"

# Cấu hình API Key
mistral_api_key = st.text_input("Nhập Mistral API Key:", type="password")

uploaded_pdf = st.file_uploader("Chọn file PDF cần chuyển đổi", type=["pdf"])

# Hàm tạo file Word bằng python-docx (Hỗ trợ chèn ảnh và text chính xác)
def generate_word_document(markdown_content, images_dict):
    doc = docx.Document()
    
    # Tách các dòng hoặc các khối để đưa vào docx
    lines = markdown_content.split("\n")
    for line in lines:
        # Kiểm tra nếu dòng chứa thẻ ảnh dạng ![alt](img_name.jpeg)
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            img_filename = os.path.basename(img_match.group(2))
            if img_filename in images_dict:
                # Lưu ảnh tạm thời ra đĩa để python-docx đọc và chèn vào docx
                temp_img_path = f"temp_{img_filename}"
                with open(temp_img_path, "wb") as img_f:
                    img_f.write(images_dict[img_filename])
                try:
                    # Chèn ảnh vào file docx với chiều rộng tối đa 5 inches
                    doc.add_picture(temp_img_path, width=Inches(4.5))
                except:
                    pass
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
        else:
            if line.strip().startswith("###"):
                doc.add_heading(line.replace("###", "").strip(), level=3)
            elif line.strip().startswith("##"):
                doc.add_heading(line.replace("##", "").strip(), level=2)
            elif line.strip().startswith("#"):
                doc.add_heading(line.replace("#", "").strip(), level=1)
            elif line.strip():
                doc.add_paragraph(line.strip())
                
    # Lưu ra file byte stream để tải về
    output_path = "output_generated.docx"
    doc.save(output_path)
    with open(output_path, "rb") as f:
        docx_data = f.read()
    if os.path.exists(output_path):
        os.remove(output_path)
    return docx_data

if uploaded_pdf and st.button("🚀 Gửi PDF lên Mistral OCR"):
    if not mistral_api_key:
        st.error("Vui lòng nhập Mistral API Key!")
    elif not MISTRAL_AVAILABLE:
        st.error("Chưa cài đặt thư viện `mistralai` trong requirements.txt!")
    else:
        with st.spinner("Đang gửi PDF lên Mistral OCR API và tạo cấu trúc..."):
            try:
                # 1. Gọi API Mistral OCR
                client = Mistral(api_key=mistral_api_key)
                file_bytes = uploaded_pdf.getvalue()
                base64_file = base64.b64encode(file_bytes).decode('utf-8')

                ocr_response = client.ocr.process(
                    document={
                        "type": "document_url",
                        "document_url": f"data:application/pdf;base64,{base64_file}"
                    },
                    model="mistral-ocr-latest",
                    include_image_base64=True,
                    include_blocks=True
                )
                
                full_markdown = ""
                images_dict = {}
                
                if hasattr(ocr_response, "pages"):
                    for idx, page in enumerate(ocr_response.pages):
                        page_md = page.markdown if hasattr(page, "markdown") else ""
                        full_markdown += f"\n\n---\n### Trang {idx+1}\n\n" + page_md
                        
                        if hasattr(page, "images") and page.images:
                            for img in page.images:
                                if hasattr(img, "id") and hasattr(img, "image_base64") and img.image_base64:
                                    img_id = img.id
                                    img_b64 = img.image_base64
                                    if "," in img_b64: img_b64 = img_b64.split(",")[1]
                                    try:
                                        img_bytes = base64.b64decode(img_b64)
                                        img_filename = f"{img_id}.jpeg"
                                        images_dict[img_filename] = img_bytes
                                    except: pass

                # 2. Tạo file Word bằng python-docx
                docx_bytes = generate_word_document(full_markdown, images_dict)

                # Lưu vào Session State
                st.session_state.preview_markdown = full_markdown
                st.session_state.images_dict = images_dict
                st.session_state.docx_bytes = docx_bytes
                st.session_state.file_name = uploaded_pdf.name.rsplit('.', 1)[0]
                
                st.success("🎉 Xử lý thành công! Vui lòng xem bản xem trước và tải file Word bên dưới.")
                
            except Exception as e:
                st.error(f"Đã xảy ra lỗi trong quá trình xử lý: {e}")

# ==========================================
# 👁️ KHUNG PREVIEW CHÍNH XÁC VỊ TRÍ ẢNH & EQUATION
# ==========================================
if st.session_state.preview_markdown:
    st.divider()
    
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("👁️ Bản xem trước Nội dung & Hình ảnh")
    with col2:
        if st.session_state.docx_bytes:
            st.download_button(
                label="📥 Tải xuống file Word (.docx)",
                data=st.session_state.docx_bytes,
                file_name=f"{st.session_state.file_name}_Converted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    # Xử lý thay thế thẻ ảnh bằng chuỗi base64 đúng tên file để hiển thị trong khung preview HTML
    raw_md = st.session_state.preview_markdown
    
    def render_preview_with_images(match):
        alt_text = match.group(1)
        img_filename = os.path.basename(match.group(2))
        if img_filename in st.session_state.images_dict:
            b64_data = base64.b64encode(st.session_state.images_dict[img_filename]).decode('utf-8')
            # Trả về thẻ HTML căn giữa đúng vị trí xuất hiện của ảnh
            return f'<div style="text-align: center; margin: 20px 0;"><img src="data:image/jpeg;base64,{b64_data}" style="max-width: 450px; border-radius: 6px; border: 1px solid #cbd5e0;" /></div>'
        return match.group(0)

    processed_html = re.sub(r'!\[(.*?)\]\((.*?)\)', render_preview_with_images, raw_md)
    
    # Chuẩn hóa công thức toán học
    processed_html = processed_html.replace(r"\(", "$").replace(r"\)", "$")
    processed_html = processed_html.replace(r"\[", "$$").replace(r"\]", "$$")
    formatted_html = processed_html.replace("\n", "<br>")

    preview_container = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/katex.min.css">
        <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/katex.min.js"></script>
        <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/contrib/auto-render.min.js" onload="renderMathInElement(document.body, {{delimiters: [{{left: '$', right: '$', display: false}}, {{left: '$$', right: '$$', display: true}}]}});"></script>
        <style>
            body {{ font-family: Arial, sans-serif; padding: 15px; background-color: #ffffff; color: #2d3748; line-height: 1.8; }}
            .preview-box {{ padding: 30px; border-radius: 10px; border: 1px solid #cbd5e0; max-height: 600px; overflow-y: auto; background: #fff; }}
            img {{ max-width: 100%; height: auto; display: block; margin: 15px auto; border-radius: 6px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
            th, td {{ border: 1px solid #cbd5e0; padding: 8px; text-align: left; }}
        </style>
    </head>
    <body>
        <div class="preview-box">
            {formatted_html}
        </div>
    </body>
    </html>
    """
    import streamlit.components.v1 as components
    components.html(preview_container, height=650, scrolling=True)